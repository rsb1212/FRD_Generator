import os

"""
SOP + RGT → FRD Generator  (with Free AI Module)
"""

from dotenv import load_dotenv
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '.env'))

from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, session
import re
import json
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
from datetime import datetime
from abc import ABC, abstractmethod

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

from ai_module import (
    compare_documents, enhance_frd, generate_ai_frd, ai_status,
    _ollama_available
)
from auth import authenticate_user

app = Flask(__name__)
app.secret_key = 'sop-rgt-frd-synth-2026-ai'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SESSION_COOKIE_HTTPONLY'] = True

app.config['LAST_SOP_TEXT'] = ""
app.config['LAST_RGT_TEXT'] = ""
app.config['LAST_FRD'] = None

def extract_docx_paragraphs(file_storage):
    if HAS_PYTHON_DOCX:
        file_storage.seek(0)
        doc = Document(file_storage)
        return [p.text for p in doc.paragraphs if p.text.strip()]
    file_storage.seek(0)
    with zipfile.ZipFile(file_storage) as zf:
        if 'word/document.xml' not in zf.namelist():
            return []
        xml_content = zf.read('word/document.xml')
    tree = ET.fromstring(xml_content)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    paragraphs = []
    for p in tree.findall('.//w:p', ns):
        texts = [t.text for t in p.findall('.//w:t', ns) if t.text]
        if texts:
            paragraphs.append(''.join(texts))
    return paragraphs

def extract_xlsx_data(file_storage):
    if HAS_OPENPYXL:
        file_storage.seek(0)
        wb = openpyxl.load_workbook(file_storage, data_only=True)
        sheets = []
        for name in wb.sheetnames:
            ws = wb[name]
            sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                sheet_rows.append([str(c) if c is not None else '' for c in row])
            sheets.append(sheet_rows)
        return sheets
    file_storage.seek(0)
    with zipfile.ZipFile(file_storage) as zf:
        shared_strings = []
        if 'xl/sharedStrings.xml' in zf.namelist():
            ss_tree = ET.fromstring(zf.read('xl/sharedStrings.xml'))
            ns = {'s': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            for si in ss_tree.findall('.//s:si', ns):
                texts = [t.text for t in si.findall('.//s:t', ns) if t.text]
                shared_strings.append(''.join(texts))
        ns = {'s': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
        sheets = []
        for sheet_file in ['xl/worksheets/sheet1.xml', 'xl/worksheets/sheet2.xml',
                           'xl/worksheets/sheet3.xml', 'xl/worksheets/sheet4.xml',
                           'xl/worksheets/sheet5.xml', 'xl/worksheets/sheet6.xml',
                           'xl/worksheets/sheet7.xml']:
            if sheet_file not in zf.namelist():
                continue
            sheet_tree = ET.fromstring(zf.read(sheet_file))
            rows = []
            for row in sheet_tree.findall('.//s:row', ns):
                cells = []
                for c in row.findall('.//s:c', ns):
                    cell_type = c.get('t', '')
                    v = c.find('s:v', ns)
                    val = ''
                    if v is not None and v.text:
                        if cell_type == 's':
                            idx = int(v.text)
                            if idx < len(shared_strings):
                                val = shared_strings[idx]
                        else:
                            val = v.text
                    cells.append(val)
                rows.append(cells)
            sheets.append(rows)
    return sheets

# ==============================================================================
# DESIGN PATTERNS & DATA STRUCTURES FOR HIGH-PERFORMANCE DOCUMENT PROCESSING
# ==============================================================================

# --- Singleton / Flyweight: Pre-compiled Regex Patterns & Frozen Lookup Sets ---
RAW_HEADER_PATTERNS = [
    r'PURPOSE:', r'SCOPE:', r'ROLES AND RESPONSIBILITY:',
    r'PROCESS METRIC:', r'PROCESS FLOW:', r'PROCEDURE:',
    r'USER CHECKLIST:', r'COMMUNICATION:', r'ESCALATION MATRIX:',
    r'REFERENCES:', r'EXCEPTION', r'KEYWORDS AND ACRONYMS:',
    r'VERSION HISTORY:', r'SIPOC:', r'SERVICE CONTINUITY:',
    r'OPUS TO NGIN MIGRATION:', r'REVIEW PROCEDURE:'
]
COMPILED_HEADER_PATTERNS = [(pat, re.compile(pat, re.IGNORECASE)) for pat in RAW_HEADER_PATTERNS]
NOTE_REGEX = re.compile(r'Note.*?(?=\n\n|\Z)', re.IGNORECASE | re.DOTALL)

DEPT_KEYWORDS = frozenset({'Branch', 'Digital', 'CCM', 'WCC', 'GRO'})
APP_TARGETS = frozenset({'OPUS', 'NGIN', 'CRM', 'PMAC', 'INSTAB', 'Cashier'})
PROD_TARGETS = frozenset({'ULIP', 'Term', 'Endowment', 'Annuity', 'Health', 'Group'})


# --- Strategy Pattern: Document Extractor Interface ---
class IDocumentExtractor(ABC):
    """Abstract Strategy interface for document data extractors."""
    @abstractmethod
    def extract(self) -> dict:
        pass


class SOPExtractor(IDocumentExtractor):
    """SOP Document Extractor Strategy using pre-compiled regexes & fast dictionary structures."""
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs
        self.sections = self._split_sections()

    def _split_sections(self):
        sections = {}
        current = 'HEADER'
        content = []
        for para in self.paragraphs:
            p = para.strip()
            matched = False
            for pat, regex in COMPILED_HEADER_PATTERNS:
                if regex.search(p):
                    if content:
                        sections[current] = '\n'.join(content)
                    current = pat.replace(':', '').strip()
                    content = [p]
                    matched = True
                    break
            if not matched:
                content.append(p)
        if content:
            sections[current] = '\n'.join(content)
        return sections

    def get_purpose(self):
        return self.sections.get('PURPOSE', '').replace('PURPOSE:', '').strip()

    def get_scope(self):
        return self.sections.get('SCOPE', '').replace('SCOPE:', '').strip()

    def get_roles(self):
        txt = self.sections.get('ROLES AND RESPONSIBILITY', '')
        roles = []
        for line in txt.split('\n'):
            if '**' in line and '---' not in line and 'ROLE' not in line:
                clean = line.replace('**', '').strip()
                if clean and len(clean) > 3:
                    roles.append(clean)
        return list(dict.fromkeys(roles))

    def get_procedure_steps(self):
        proc = self.sections.get('PROCEDURE', '')
        steps = []
        current = {}
        for line in proc.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('**') and '**' in line[2:]:
                if current:
                    steps.append(current)
                name = line.replace('**', '').strip()
                current = {'name': name, 'tasks': [], 'dept': ''}
            elif line.startswith(('-', '•')):
                current.setdefault('tasks', []).append(line.lstrip('- •').strip())
            elif any(k in line for k in DEPT_KEYWORDS):
                current['dept'] = line.replace('**', '').strip()
        if current:
            steps.append(current)
        return steps

    def get_business_rules(self):
        rules = []
        for sec_name in ('SIPOC', 'PROCEDURE'):
            text = self.sections.get(sec_name, '')
            for note in NOTE_REGEX.findall(text):
                for ln in note.split('\n'):
                    cl = ln.strip().lstrip('0123456789.').lstrip('-').strip()
                    if cl and 'Note' not in cl and len(cl) > 10:
                        rules.append(cl)
        return list(dict.fromkeys(rules))

    def get_escalation(self):
        return self.sections.get('ESCALATION MATRIX', '')

    def extract(self) -> dict:
        return {
            'purpose': self.get_purpose(),
            'scope': self.get_scope(),
            'roles': self.get_roles(),
            'procedure_steps': self.get_procedure_steps(),
            'business_rules': self.get_business_rules(),
            'escalation': self.get_escalation()
        }


class RGTExtractor(IDocumentExtractor):
    """Optimized RGT Extractor Strategy — single-pass scan with cached results."""
    def __init__(self, sheets):
        self.sheets = sheets
        self.flat = ' '.join(' '.join(str(c) for c in row if c) for sheet in sheets for row in sheet)
        self._cache = {}
        self._extract_all()

    def _extract_all(self):
        """Single pass over all sheets to extract everything at once — avoids repeated scans."""
        deltas = []
        scenarios = []
        apps = []
        prods = []
        call_summary = None
        reqs = []
        capture_req = False

        for sheet in self.sheets:
            in_delta = False
            in_test = False
            for row in sheet:
                if not row:
                    continue
                txt = ' '.join(str(c) for c in row if c)
                first = str(row[0]).strip()

                # --- Call Summary ---
                if call_summary is None and 'frequency change' in txt.lower() and len(txt) > 20:
                    if any(k in txt.lower() for k in ['enhancement', 'process', 'detailed']):
                        call_summary = txt

                # --- Business Requirements ---
                if 'PART B' in txt or 'Detailed Requirement' in txt:
                    capture_req = True
                if 'PART C' in txt:
                    capture_req = False
                if capture_req and len(txt) > 20 and any(k in txt for k in ['Invoice', 'EFT', '•', 'handling']):
                    reqs.append(txt)

                # --- Delta Features ---
                if 'Delta feature' in txt or (len(row) >= 3 and first == 'Functionality'):
                    in_delta = True
                    continue
                if in_delta and len(row) >= 3:
                    f, e, r = str(row[0]), str(row[1]), str(row[2])
                    if f and f != 'Functionality':
                        deltas.append({'functionality': f, 'expected': e, 'requirement': r})
                elif in_delta and len(row) < 3:
                    in_delta = False

                # --- Test Scenarios ---
                if 'Test Scenario' in txt or 'Revised_Test' in txt:
                    in_test = True
                    continue
                if in_test and len(row) >= 3:
                    scenarios.append({
                        'scenario': str(row[0]),
                        'module': str(row[1]) if len(row) > 1 else '',
                        'condition': str(row[2]) if len(row) > 2 else '',
                        'expected': str(row[3]) if len(row) > 3 else ''
                    })
                elif in_test and len(row) < 3:
                    in_test = False

                # --- Applications Impacted ---
                if any(t in first for t in APP_TARGETS):
                    if any(str(c).strip().lower() == 'yes' for c in row[1:]):
                        apps.append(first)

                # --- Products Impacted ---
                if any(t in first for t in PROD_TARGETS):
                    if len(row) > 1 and str(row[1]).strip().lower() == 'yes':
                        prods.append(first)

        self._cache['call_summary'] = call_summary or 'Frequency Change Process Enhancement'
        self._cache['business_requirement'] = reqs
        self._cache['delta_features'] = deltas
        self._cache['test_scenarios'] = scenarios
        self._cache['applications_impacted'] = list(dict.fromkeys(apps))
        self._cache['products_impacted'] = list(dict.fromkeys(prods))

    def get_call_summary(self):
        return self._cache['call_summary']
    def get_business_requirement(self):
        return self._cache['business_requirement']
    def get_delta_features(self):
        return self._cache['delta_features']
    def get_test_scenarios(self):
        return self._cache['test_scenarios']
    def get_applications_impacted(self):
        return self._cache['applications_impacted']
    def get_products_impacted(self):
        return self._cache['products_impacted']
    def get_priority(self):
        if 'Must Have' in self.flat:
            return 'High (Must Have)'
        if 'Good to Have' in self.flat:
            return 'Medium (Good to Have)'
        return 'Medium'
    def get_process_change(self):
        return 'Yes — Major Change' if ('Process Change' in self.flat and 'Yes' in self.flat) else 'No'

    def extract(self) -> dict:
        return self._cache


# --- Factory Pattern: Extractor Factory ---
class ExtractorFactory:
    """Factory to instantiate the appropriate document extractor strategy."""
    @staticmethod
    def create_sop_extractor(paragraphs) -> SOPExtractor:
        return SOPExtractor(paragraphs)

    @staticmethod
    def create_rgt_extractor(sheets) -> RGTExtractor:
        return RGTExtractor(sheets)


def generate_frd(sop, rgt):
    call_id = f"CR-FREQ-{datetime.now():%Y-%m%d}"
    summary = rgt.get_call_summary()
    business_requirement = rgt.get_business_requirement()
    if not business_requirement:
        business_requirement = [
            'Enable policy servicing teams to update premium frequency efficiently without manual rework.',
            'Ensure the change is reflected across downstream systems with consistent business rules.'
        ]
    core_steps = [
        'Validate existing policy and frequency setup.',
        'Capture the requested change and related exceptions.',
        'Trigger downstream updates in the servicing and policy administration platforms.',
        'Generate confirmation communication and maintain audit trail.'
    ]
    stakeholders = list(dict.fromkeys([
        'Policy Servicing (PS)', 'Branch Operations', 'CCM', 'IT Development',
        'Quality Assurance (QA)', 'Business Analysis', 'Digital Servicing', 'Logistics'
    ]))
    rules = sop.get_business_rules()
    if not rules:
        rules = [
            'All frequency changes must be logged with request reference and effective date.',
            'Any exception or policy-specific condition must be routed to business approvers before processing.'
        ]
    steps = []
    n = 1
    for proc in sop.get_procedure_steps():
        for task in proc.get('tasks', []):
            if len(task) > 15:
                steps.append({'step': n, 'description': task, 'actor': proc.get('dept', 'System')})
                n += 1
    if not steps:
        def infer_actor(text):
            lower_text = text.lower()
            if any(word in lower_text for word in ('customer', 'policyholder', 'self-service')):
                return 'Customer'
            if any(word in lower_text for word in ('kyc', 'branch', 'document')):
                return 'Branch/WCC/GRO'
            if any(word in lower_text for word in ('crm', 'service request', 'sr raised')):
                return 'Touchpoint User'
            if any(word in lower_text for word in ('communication', 'notification', 'email', 'sms')):
                return 'CCM'
            if any(word in lower_text for word in ('system', 'opus', 'ngin', 'scheduler', 'invoice', 'eft')):
                return 'System'
            return 'Business Servicing'

        candidate_descriptions = list(business_requirement)
        candidate_descriptions.extend(
            feature.get('functionality', '')
            for feature in rgt.get_delta_features()
            if feature.get('functionality')
        )
        seen_descriptions = set()
        for candidate in candidate_descriptions:
            fragments = re.split(r'(?<=[.!?])\s+|[;\n]+', str(candidate).strip())
            for description in fragments:
                description = re.sub(r'^[\s\-•\d.)]+', '', description).strip()
                key = description.lower()
                if len(description) > 15 and key not in seen_descriptions:
                    seen_descriptions.add(key)
                    steps.append({
                        'step': len(steps) + 1,
                        'description': description[:180],
                        'actor': infer_actor(description)
                    })
                if len(steps) >= 12:
                    break
            if len(steps) >= 12:
                break

        if not steps:
            steps = [
                {'step': 1, 'description': 'Customer approaches touchpoint for frequency change', 'actor': 'Customer'},
                {'step': 2, 'description': 'System validates policy eligibility and status', 'actor': 'System'},
                {'step': 3, 'description': 'KYC documents collected and verified', 'actor': 'Branch/WCC/GRO'},
                {'step': 4, 'description': 'SR raised in CRM with new frequency details', 'actor': 'Touchpoint User'},
                {'step': 5, 'description': 'System processes change (OPUS/NGIN/Scheduler)', 'actor': 'System'},
                {'step': 6, 'description': 'Auto-cancels/regenerates invoice if applicable', 'actor': 'System'},
                {'step': 7, 'description': 'Communication triggered to policyholder', 'actor': 'CCM'},
            ]
    deps = [
        'OPUS Core System', 'NGIN Policy Administration', 'CRM (iAhead)',
        'Customer Portal / Life Assist App', 'CCM', 'GPA',
        'Scheduler / Reverse Feed', 'Payment Gateway'
    ]
    for a in rgt.get_applications_impacted():
        if a not in [d.split()[0] for d in deps]:
            deps.append(f"{a} Application")
    return {
        'call_id': call_id,
        'call_summary': summary,
        'objective': sop.get_purpose() or 'Modernize the frequency change process to ensure consistent, auditable, and compliant servicing outcomes.',
        'business_requirement': business_requirement,
        'applicable_platform': ', '.join(rgt.get_applications_impacted()) or 'OPUS, NGIN, CRM',
        'priority': rgt.get_priority(),
        'stakeholders': stakeholders,
        'current_process': 'Manual SR raising with backend processing. OPUS: CRM → OPUS → Validate → Rate → Save. NGIN: CRM → Scheduler → Reverse Feed. Exceptions handled by PS via ITHD.',
        'future_process': 'Enhanced automated frequency change with invoice auto-cancellation, EFT mode handling without DC activation, digital self-service, automated communication, and configurator-driven rules.',
        'functional_scope': [
            'Capture and validate frequency change requests',
            'Apply business rules and exception handling',
            'Trigger downstream updates in policy and servicing systems',
            'Provide audit trail and communication updates for policyholders'
        ],
        'non_functional_requirements': [
            'The solution must support role-based access for servicing and operations teams.',
            'All changes must be traceable and available for audit review.',
            'The process must provide measurable turnaround times and error handling for failed downstream updates.'
        ],
        'acceptance_criteria': [
            'A valid request completes end-to-end without manual intervention.',
            'Invalid or exception scenarios are flagged for business review.',
            'System-generated communication and documentation are available for the case record.'
        ],
        'dependencies': deps,
        'out_of_scope': [
            'Invest Protect Goal II (explicitly excluded)',
            'Finance/Accounting changes (no financial impact)',
            'Communication format changes (not required)',
            'Policy surrender processing', 'Claim processing integration'
        ],
        'business_rules': rules,
        'process_json': {
            'process_name': 'Extracted Process',
            'nodes': [{'id': f'N{i+1}', 'actor': s.get('actor', 'System') if isinstance(s, dict) else 'System', 'type': 'process', 'label': s.get('description', s) if isinstance(s, dict) else str(s)[:50]} for i, s in enumerate(steps)] if steps else [
                {'id': 'N1', 'actor': 'Branch/Servicing Team', 'type': 'process', 'label': core_steps[0]},
                {'id': 'N2', 'actor': 'Business Analyst', 'type': 'process', 'label': core_steps[1]},
                {'id': 'N3', 'actor': 'System / Integration Layer', 'type': 'process', 'label': core_steps[2]},
                {'id': 'N4', 'actor': 'CCM / Policy Servicing', 'type': 'process', 'label': core_steps[3]}
            ],
            'connections': [{'from': f'N{i+1}', 'to': f'N{i+2}'} for i in range(len(steps) - 1)] if steps else [
                {'from': 'N1', 'to': 'N2'},
                {'from': 'N2', 'to': 'N3'},
                {'from': 'N3', 'to': 'N4'}
            ]
        },
        'test_scenarios': rgt.get_test_scenarios(),
        'delta_features': rgt.get_delta_features(),
        'products_impacted': rgt.get_products_impacted(),
        'process_change': rgt.get_process_change(),
        'escalation_matrix': sop.get_escalation(),
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'ai_enhanced': False,
        'assumptions': [],
        'compliance_notes': 'The solution must align with policy servicing controls, data privacy expectations, and audit requirements for change management.',
        'risk_flags': []
    }


def login_required(view):
    def wrapped(*args, **kwargs):
        if not session.get('logged_in'):
            flash('Please sign in to continue.')
            return redirect(url_for('login'))
        return view(*args, **kwargs)
    wrapped.__name__ = view.__name__
    return wrapped

def dedupe_list(items):
    """Deduplicate list of strings or objects while preserving original order."""
    if not items:
        return []
    seen = set()
    result = []
    for item in items:
        if isinstance(item, str):
            clean = item.strip()
            key = clean.lower()
            if clean and key not in seen:
                seen.add(key)
                result.append(clean)
        elif isinstance(item, dict):
            key = json.dumps(item, sort_keys=True).lower()
            if key not in seen:
                seen.add(key)
                result.append(item)
        else:
            if item not in seen:
                seen.add(item)
                result.append(item)
    return result

def generate_mermaid_diagram(process_json):
    """Generate Mermaid.js flowchart string from process_json."""
    if not process_json or not isinstance(process_json, dict) or 'nodes' not in process_json:
        return ""
    lines = [
        "flowchart LR",
        "    classDef startend fill:#dae8fc,stroke:#6c8ebf,stroke-width:2px;",
        "    classDef process fill:#d5e8d4,stroke:#82b366,stroke-width:2px;",
        "    classDef decision fill:#ffe6cc,stroke:#d79b00,stroke-width:2px;"
    ]
    
    nodes = process_json.get('nodes', [])
    connections = process_json.get('connections', [])
    
    # Group nodes by actor (swimlanes)
    actors = {}
    for n in nodes:
        a = n.get('actor', 'System').replace('"', '').replace("'", "")
        if a not in actors:
            actors[a] = []
        actors[a].append(n)
        
    for actor, actor_nodes in actors.items():
        lines.append(f'    subgraph {actor.replace(" ", "_")} ["{actor}"]')
        for n in actor_nodes:
            nid = n.get('id')
            label = n.get('label', '').replace('"', '').replace("'", "")
            ntype = n.get('type', 'process').lower()
            if ntype in ['start', 'end']:
                lines.append(f'        {nid}(("{label}")):::startend')
            elif ntype == 'decision':
                lines.append(f'        {nid}{{"{label}"}}:::decision')
            else:
                lines.append(f'        {nid}["{label}"]:::process')
        lines.append('    end')
        
    for c in connections:
        src = c.get('from')
        dst = c.get('to')
        lbl = c.get('label', '').replace('"', '')
        if src and dst:
            if lbl:
                lines.append(f'    {src} -- "{lbl}" --> {dst}')
            else:
                lines.append(f'    {src} --> {dst}')
                
    return "\n".join(lines)


def generate_drawio_xml(process_json):
    """Generate a Draw.io swimlane diagram from process_json."""
    import html
    if not process_json or not isinstance(process_json, dict) or 'nodes' not in process_json:
        return ""

    nodes = process_json.get('nodes', [])
    connections = process_json.get('connections', [])
    
    if not nodes:
        return ""

    actors = []
    for n in nodes:
        actor = str(n.get('actor', 'System')).strip() or 'System'
        if actor not in actors:
            actors.append(actor)

    lane_width = 240
    # Calculate lane height based on the maximum number of nodes in any single actor's lane
    max_nodes_in_lane = max((sum(1 for n in nodes if str(n.get('actor', 'System')).strip() == actor) for actor in actors), default=0)
    lane_height = max(180, max_nodes_in_lane * 120 + 60)
    node_width = 200
    node_height = 60
    lane_x = {actor: index * lane_width for index, actor in enumerate(actors)}
    xml = [
        '<mxfile host="app.diagrams.net" version="24.4.10">',
        '  <diagram id="ProcessFlow" name="Process Flow">',
        '    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="827" pageHeight="1169" math="0" shadow="0">',
        '      <root>',
        '        <mxCell id="0" />',
        '        <mxCell id="1" parent="0" />'
    ]

    # Swimlanes
    next_id = 2
    lane_ids = {}
    for actor in actors:
        lane_id = str(next_id)
        lane_ids[actor] = lane_id
        next_id += 1
        xml.append(f'        <mxCell id="{lane_id}" value="{html.escape(actor)}" style="swimlane;whiteSpace=wrap;html=1;" vertex="1" parent="1">')
        xml.append(f'          <mxGeometry x="{lane_x[actor]}" y="0" width="{lane_width}" height="{lane_height}" as="geometry" />')
        xml.append('        </mxCell>')

    # Nodes
    node_ids = {}
    row_by_actor = {actor: 0 for actor in actors}
    for n in nodes:
        actor = str(n.get('actor', 'System')).strip() or 'System'
        label = str(n.get('label', '')).strip()
        ntype = str(n.get('type', 'process')).strip().lower()
        nid = str(n.get('id', ''))
        
        lower_label = label.lower()
        if ntype == 'start' or 'start' in lower_label or 'begin' in lower_label:
            style = 'ellipse;whiteSpace=wrap;html=1;fillColor=#dae8fc;'
        elif ntype == 'decision' or '?' in lower_label:
            style = 'rhombus;whiteSpace=wrap;html=1;fillColor=#ffe6cc;'
        elif ntype == 'end' or 'finish' in lower_label or 'complete' in lower_label:
            style = 'ellipse;whiteSpace=wrap;html=1;fillColor=#dae8fc;'
        else:
            style = 'rounded=1;whiteSpace=wrap;html=1;fillColor=#d5e8d4;'

        row = row_by_actor[actor]
        row_by_actor[actor] += 1
        mx_id = str(next_id)
        node_ids[nid] = mx_id
        next_id += 1
        
        xml.append(f'        <mxCell id="{mx_id}" value="{html.escape(label)}" style="{style}" vertex="1" parent="{lane_ids[actor]}">')
        xml.append(f'          <mxGeometry x="20" y="{40 + row * 120}" width="{node_width}" height="{node_height}" as="geometry" />')
        xml.append('        </mxCell>')

    # Edges
    for c in connections:
        src = str(c.get('from', ''))
        dst = str(c.get('to', ''))
        edge_label = str(c.get('label', '')).strip()
        
        if src in node_ids and dst in node_ids:
            edge_id = str(next_id)
            next_id += 1
            if edge_label:
                xml.append(f'        <mxCell id="{edge_id}" value="{html.escape(edge_label)}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;" edge="1" parent="1" source="{node_ids[src]}" target="{node_ids[dst]}">')
            else:
                xml.append(f'        <mxCell id="{edge_id}" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;endArrow=block;" edge="1" parent="1" source="{node_ids[src]}" target="{node_ids[dst]}">')
            xml.append('          <mxGeometry relative="1" as="geometry" />')
            xml.append('        </mxCell>')

    xml.extend([
        '      </root>',
        '    </mxGraphModel>',
        '  </diagram>',
        '</mxfile>'
    ])
    return "\n".join(xml)


def add_flowchart_to_docx(doc, process_steps):
    """Add a visual process flow diagram to a Word document using styled tables."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    if not process_steps:
        return

    def _is_actor_step(actor):
        return any(a in actor.lower() for a in ['customer', 'user', 'branch', 'ps', 'gro', 'wcc', 'agent', 'ccm', 'touchpoint'])

    def _set_cell_bg(cell, hex_color):
        """Set cell background color."""
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): hex_color
        })
        shading_elm.append(shading)

    def _set_cell_borders(cell, color='CCCCCC'):
        """Set rounded-look borders on a cell."""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge in ('top', 'left', 'bottom', 'right'):
            border = borders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): 'single',
                qn('w:sz'): '8',
                qn('w:space'): '0',
                qn('w:color'): color
            })
            borders.append(border)
        tcPr.append(borders)

    # Build the flowchart: each step is a row, with an arrow row between steps
    total_rows = len(process_steps) * 2 - 1  # step + arrow rows (no arrow after last)
    table = doc.add_table(rows=total_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all default borders from the table
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders_el = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = borders_el.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders_el.append(border)
    tblPr.append(borders_el)

    row_idx = 0
    for i, step in enumerate(process_steps):
        if not isinstance(step, dict):
            step = {'step': i + 1, 'actor': 'System', 'description': str(step)}

        actor = step.get('actor', 'System')
        desc = step.get('description', '')
        step_num = step.get('step', i + 1)
        is_actor = _is_actor_step(actor)

        # Step row — use the center cell only
        row_cells = table.rows[row_idx].cells
        # Leave left and right cells empty (spacers)
        center_cell = row_cells[1]

        # Color: purple-blue for actors, teal-green for system
        bg_color = '667EEA' if is_actor else '00B09B'
        border_color = '5A6FD1' if is_actor else '009688'

        _set_cell_bg(center_cell, bg_color)
        _set_cell_borders(center_cell, border_color)

        # Step number + actor line
        p_actor = center_cell.paragraphs[0]
        p_actor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_actor = p_actor.add_run(f"Step {step_num}: {actor}")
        run_actor.bold = True
        run_actor.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_actor.font.size = Pt(10)

        # Description line
        p_desc = center_cell.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        short_desc = (desc[:60] + '...') if len(desc) > 60 else desc
        run_desc = p_desc.add_run(short_desc)
        run_desc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_desc.font.size = Pt(9)

        # Make spacer cells borderless and transparent
        for spacer_cell in [row_cells[0], row_cells[2]]:
            _set_cell_bg(spacer_cell, 'FFFFFF')
            # Clear borders
            tc = spacer_cell._element
            tcPr = tc.get_or_add_tcPr()
            sb = tcPr.makeelement(qn('w:tcBorders'), {})
            for edge in ('top', 'left', 'bottom', 'right'):
                b = sb.makeelement(qn(f'w:{edge}'), {
                    qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                })
                sb.append(b)
            tcPr.append(sb)

        row_idx += 1

        # Arrow row (between steps, not after the last one)
        if i < len(process_steps) - 1:
            arrow_cells = table.rows[row_idx].cells
            arrow_cell = arrow_cells[1]
            p_arrow = arrow_cell.paragraphs[0]
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_arrow = p_arrow.add_run('▼')
            run_arrow.font.size = Pt(16)
            run_arrow.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
            run_arrow.bold = True

            # Make all arrow row cells borderless
            for ac in arrow_cells:
                _set_cell_bg(ac, 'FFFFFF')
                tc = ac._element
                tcPr = tc.get_or_add_tcPr()
                sb = tcPr.makeelement(qn('w:tcBorders'), {})
                for edge in ('top', 'left', 'bottom', 'right'):
                    b = sb.makeelement(qn(f'w:{edge}'), {
                        qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                    })
                    sb.append(b)
                tcPr.append(sb)

            row_idx += 1

    # Add a legend below the diagram
    legend_p = doc.add_paragraph()
    legend_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend_p.paragraph_format.space_before = Pt(8)
    run_legend1 = legend_p.add_run('■ ')
    run_legend1.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
    run_legend1.font.size = Pt(10)
    run_label1 = legend_p.add_run('Actor/User Step    ')
    run_label1.font.size = Pt(9)
    run_legend2 = legend_p.add_run('■ ')
    run_legend2.font.color.rgb = RGBColor(0x00, 0xB0, 0x9B)
    run_legend2.font.size = Pt(10)
    run_label2 = legend_p.add_run('System Step')
    run_label2.font.size = Pt(9)

def clean_and_dedupe_frd(frd):
    """Deduplicate all list attributes in FRD and generate flow diagram data."""
    if not frd or not isinstance(frd, dict):
        return frd
    list_fields = [
        'business_requirement', 'business_rules', 'stakeholders',
        'dependencies', 'out_of_scope', 'assumptions', 'risk_flags',
        'functional_scope', 'non_functional_requirements', 'acceptance_criteria',
        'products_impacted'
    ]
    for field in list_fields:
        if field in frd and isinstance(frd[field], list):
            frd[field] = dedupe_list(frd[field])
            
    if 'process_json' in frd and isinstance(frd['process_json'], dict):
        pass # Trust AI structured JSON
    elif 'process_steps' in frd and isinstance(frd['process_steps'], list):
        # Backward compatibility conversion to process_json
        nodes = []
        connections = []
        seen_steps = set()
        clean_steps = []
        for step in frd['process_steps']:
            if isinstance(step, dict):
                desc = step.get('description', '').strip().lower()
                if desc and desc not in seen_steps:
                    seen_steps.add(desc)
                    clean_steps.append(step)
            elif isinstance(step, str):
                desc = step.strip().lower()
                if desc and desc not in seen_steps:
                    seen_steps.add(desc)
                    clean_steps.append({'actor': 'System', 'description': step})
                    
        for i, s in enumerate(clean_steps, 1):
            nodes.append({'id': f'N{i}', 'actor': s.get('actor', 'System'), 'type': 'process', 'label': s.get('description', f'Step {i}')})
            if i < len(clean_steps):
                connections.append({'from': f'N{i}', 'to': f'N{i+1}'})
        
        frd['process_json'] = {
            'process_name': 'Legacy Converted Flow',
            'nodes': nodes,
            'connections': connections
        }

    if 'delta_features' in frd and isinstance(frd['delta_features'], list):
        seen_deltas = set()
        clean_deltas = []
        for d in frd['delta_features']:
            if isinstance(d, dict):
                func = d.get('functionality', '').strip().lower()
                if func and func not in seen_deltas:
                    seen_deltas.add(func)
                    clean_deltas.append(d)
        frd['delta_features'] = clean_deltas

    if 'test_scenarios' in frd and isinstance(frd['test_scenarios'], list):
        seen_tests = set()
        clean_tests = []
        for t in frd['test_scenarios']:
            if isinstance(t, dict):
                scen = t.get('scenario', '').strip().lower()
                if scen and scen not in seen_tests:
                    seen_tests.add(scen)
                    clean_tests.append(t)
        frd['test_scenarios'] = clean_tests

    # Generate Mermaid diagram & process flow step chain
    pjson = frd.get('process_json', {})
    frd['mermaid_diagram'] = generate_mermaid_diagram(pjson)
    frd['drawio_xml'] = generate_drawio_xml(pjson)
    flow_nodes = [f"[{n.get('actor', 'System')}: {n.get('label', '')[:35]}]" for n in pjson.get('nodes', [])]
    frd['process_flow_chain'] = " -> ".join(flow_nodes)

    return frd

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')
        if authenticate_user(username, password):
            session['logged_in'] = True
            session['username'] = username
            flash('Welcome back.')
            return redirect(url_for('index'))
        flash('Invalid username or password.')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been signed out.')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    return render_template('index.html', ai_status=ai_status())

@app.route('/ai-status')
@login_required
def ai_status_route():
    return jsonify(ai_status())

@app.route('/generate', methods=['POST'])
@login_required
def generate():
    rgt_f = request.files.get('rgt_file')
    ai_module = request.form.get('ai_module', 'gemini')
    if not rgt_f or rgt_f.filename == '':
        flash('Please upload the RGT file.')
        return redirect(url_for('index'))
    if not rgt_f.filename.endswith('.xlsx'):
        flash('RGT must be a .xlsx file.')
        return redirect(url_for('index'))

    try:
        rgt_sheets = extract_xlsx_data(rgt_f)
        sop = ExtractorFactory.create_sop_extractor([])
        rgt = ExtractorFactory.create_rgt_extractor(rgt_sheets)
        sop_text = ""
        rgt_text = '\n'.join(' '.join(str(c) for c in row if c) for sheet in rgt_sheets for row in sheet)
        app.config['LAST_SOP_TEXT'] = sop_text
        app.config['LAST_RGT_TEXT'] = rgt_text
        
        frd = generate_ai_frd(sop_text, rgt_text, provider=ai_module)
        if 'error' in frd:
            flash(f"AI Generation failed: {frd['error']}. Falling back to standard.")
            frd = generate_frd(sop, rgt)
        else:
            frd.setdefault('test_scenarios', rgt.get_test_scenarios())
            frd.setdefault('delta_features', rgt.get_delta_features())
            frd.setdefault('products_impacted', rgt.get_products_impacted())
            frd.setdefault('process_change', rgt.get_process_change())
            frd.setdefault('escalation_matrix', sop.get_escalation())
            frd.setdefault('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        # Clean and deduplicate data
        frd = clean_and_dedupe_frd(frd)
        app.config['LAST_FRD'] = frd
        return render_template('frd.html', frd=frd)
    except Exception as e:
        flash(f'Processing error: {str(e)}')
        return redirect(url_for('index'))

@app.route('/ai-compare')
@login_required
def ai_compare():
    flash('AI Compare is disabled since SOP upload is removed.')
    return redirect(url_for('index'))


@app.route('/download')
@login_required
def download():
    frd = app.config.get('LAST_FRD')
    if not frd:
        flash('Generate an FRD first.')
        return redirect(url_for('index'))

    if HAS_PYTHON_DOCX:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Document Title
        title = doc.add_heading('Functional Requirement Document', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.add_run(f"Call ID: {frd.get('call_id', '')}\n").bold = True
        p.add_run(f"Call Summary: {frd.get('call_summary', '')}\n")
        p.add_run(f"Priority: {frd.get('priority', '')}\n")
        p.add_run(f"Applicable Platform: {frd.get('applicable_platform', '')}\n")
        p.add_run(f"Process Change: {frd.get('process_change', 'No')}\n")
        p.add_run(f"Generated At: {frd.get('generated_at', '')}\n")
        if frd.get('ai_enhanced'):
            doc.add_paragraph("⚡ AI-Enhanced FRD", style='Heading 3')

        # 1. Objective
        doc.add_heading('1. Objective', level=1)
        doc.add_paragraph(frd.get('objective', ''))

        # 2. Products Impacted
        if frd.get('products_impacted'):
            doc.add_heading('2. Products Impacted', level=1)
            doc.add_paragraph(', '.join(frd.get('products_impacted', [])))

        # 3. Business Requirements
        doc.add_heading('3. Business Requirements', level=1)
        for r in frd.get('business_requirement', []):
            doc.add_paragraph(r, style='List Bullet')

        # 4. Current vs Future Process Comparison
        doc.add_heading('4. Process Transformation (Current vs Future)', level=1)
        doc.add_heading('Current Process:', level=2)
        doc.add_paragraph(frd.get('current_process', ''))
        doc.add_heading('Future Process:', level=2)
        doc.add_paragraph(frd.get('future_process', ''))

        # 5. Process Flow Diagram & Flow Changes
        doc.add_heading('5. Process Flow Diagram', level=1)

        # Add visual flowchart diagram
        if frd.get('process_steps'):
            doc.add_heading('Visual Process Flow:', level=2)
            add_flowchart_to_docx(doc, frd['process_steps'])

            # Also add the detailed step table below the diagram
            doc.add_heading('Detailed Step Breakdown:', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Step #'
            hdr_cells[1].text = 'Actor / System'
            hdr_cells[2].text = 'Process Action Description'
            for cell in hdr_cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True

            for s in frd.get('process_steps', []):
                if isinstance(s, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(s.get('step', ''))
                    row_cells[1].text = str(s.get('actor', ''))
                    row_cells[2].text = str(s.get('description', ''))

        # 6. Stakeholders
        if frd.get('stakeholders'):
            doc.add_heading('6. Stakeholders', level=1)
            for s in frd.get('stakeholders', []):
                doc.add_paragraph(s, style='List Bullet')

        # 7. Business Rules
        if frd.get('business_rules'):
            doc.add_heading('7. Business Rules', level=1)
            for i, r in enumerate(frd.get('business_rules', []), 1):
                doc.add_paragraph(f"{i}. {r}", style='List Number')

        # 8. Functional Scope
        if frd.get('functional_scope'):
            doc.add_heading('8. Functional Scope', level=1)
            for item in frd.get('functional_scope', []):
                doc.add_paragraph(item, style='List Bullet')

        # 9. Non-Functional Requirements
        if frd.get('non_functional_requirements'):
            doc.add_heading('9. Non-Functional Requirements', level=1)
            for item in frd.get('non_functional_requirements', []):
                doc.add_paragraph(item, style='List Bullet')

        # 10. Acceptance Criteria
        if frd.get('acceptance_criteria'):
            doc.add_heading('10. Acceptance Criteria', level=1)
            for item in frd.get('acceptance_criteria', []):
                doc.add_paragraph(item, style='List Bullet')

        # 11. Delta Features
        if frd.get('delta_features'):
            doc.add_heading('11. Delta Features', level=1)
            for d in frd.get('delta_features', []):
                if isinstance(d, dict):
                    doc.add_paragraph(f"Functionality: {d.get('functionality', '')}\nExpected: {d.get('expected', '')}", style='List Bullet')

        # 12. Test Scenarios
        if frd.get('test_scenarios'):
            doc.add_heading('12. Test Scenarios', level=1)
            for t in frd.get('test_scenarios', []):
                if isinstance(t, dict):
                    doc.add_paragraph(f"Scenario: {t.get('scenario', '')}\nModule: {t.get('module', '')} | Condition: {t.get('condition', '')} | Expected: {t.get('expected', '')}", style='List Bullet')

        # 13. Dependencies
        if frd.get('dependencies'):
            doc.add_heading('13. Dependencies', level=1)
            for d in frd.get('dependencies', []):
                doc.add_paragraph(d, style='List Bullet')

        # 14. Out of Scope
        if frd.get('out_of_scope'):
            doc.add_heading('14. Out of Scope', level=1)
            for o in frd.get('out_of_scope', []):
                doc.add_paragraph(o, style='List Bullet')

        # 15. Assumptions
        if frd.get('assumptions'):
            doc.add_heading('15. Assumptions', level=1)
            for a in frd.get('assumptions', []):
                doc.add_paragraph(a, style='List Bullet')

        # 16. Compliance Notes
        if frd.get('compliance_notes'):
            doc.add_heading('16. Compliance Notes', level=1)
            doc.add_paragraph(frd.get('compliance_notes', ''))

        # 17. Risk Flags
        if frd.get('risk_flags'):
            doc.add_heading('17. Risk Flags', level=1)
            for rf in frd.get('risk_flags', []):
                doc.add_paragraph(rf, style='List Bullet')

        # 18. Escalation Matrix
        if frd.get('escalation_matrix'):
            doc.add_heading('18. Escalation Matrix', level=1)
            doc.add_paragraph(frd.get('escalation_matrix', ''))

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=f"FRD_{frd.get('call_id', 'CR-FREQ')}.docx")

    html = render_template('frd.html', frd=frd)
    buf = BytesIO(html.encode('utf-8'))
    buf.seek(0)
    return send_file(buf, mimetype='text/html', as_attachment=True,
                     download_name=f"FRD_{frd.get('call_id', 'CR-FREQ')}.html")

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    app.run(debug=True, host='0.0.0.0', port=5000)

